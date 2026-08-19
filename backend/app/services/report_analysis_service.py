"""
Medical report analysis pipeline.

This service:
- extracts text from uploaded PDF/DOCX/TXT reports
- falls back to OCR for scanned PDFs
- parses report facts and lab findings
- generates deterministic retrieval queries
- retrieves trusted medical evidence through a local vector index
- optionally calls Gemini for grounded JSON generation
- persists the full analysis back to the report record
"""
from __future__ import annotations

import json
import asyncio
import logging
import math
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Optional

import pdfplumber
import pytesseract
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

try:  # pragma: no cover - optional dependency guard
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover - fallback path
    fitz = None

try:  # pragma: no cover - optional dependency guard
    import docx  # type: ignore
except Exception:  # pragma: no cover - fallback path
    docx = None

from app.core.config import settings
from app.models import MedicalReport, ReportStatus
from app.rag.embeddings_manager import EmbeddingsManager
from app.rag.retriever import Retriever
from app.rag.vector_store_manager import VectorStoreManager
from app.schemas import (
    DetectedCondition,
    EvidenceSource,
    MedicalTerm,
    PatientInfo,
    ReportAnalysisDetail,
    ReportFinding,
)
from app.utils.file_handler import FileHandler
from app.utils.logger import get_logger

logger = get_logger("report_analysis_service")

try:  # pragma: no cover - optional SDK guard
    from google import genai
    from google.genai import types
except Exception:  # pragma: no cover - fallback path
    genai = None
    types = None


@dataclass(frozen=True)
class TrustedSourceDocument:
    citation_id: str
    title: str
    organization: str
    year: int
    source_type: str
    url: str
    content: str


TRUSTED_SOURCE_LIBRARY: list[TrustedSourceDocument] = [
    TrustedSourceDocument(
        citation_id="acc_aha_cholesterol_2018",
        title="2018 Guideline on the Management of Blood Cholesterol",
        organization="American College of Cardiology / American Heart Association",
        year=2018,
        source_type="guideline",
        url="https://www.ahajournals.org/doi/10.1161/CIR.0000000000000625",
        content=(
            "LDL cholesterol is a major cardiovascular risk factor. Elevated LDL, "
            "low HDL, and elevated triglycerides support a pattern of dyslipidemia. "
            "Lifestyle change is first-line for many patients, with statin therapy "
            "considered based on overall cardiovascular risk and persistent elevation."
        ),
    ),
    TrustedSourceDocument(
        citation_id="ada_standards_2024",
        title="Standards of Care in Diabetes - 2024",
        organization="American Diabetes Association",
        year=2024,
        source_type="guideline",
        url="https://diabetesjournals.org/care/issue/47/Supplement_1",
        content=(
            "Fasting plasma glucose between 100 and 125 mg/dL and HbA1c between "
            "5.7% and 6.4% are consistent with prediabetes. Fasting glucose at or "
            "above 126 mg/dL or HbA1c at or above 6.5% is consistent with diabetes "
            "and requires clinical correlation."
        ),
    ),
    TrustedSourceDocument(
        citation_id="cdc_prediabetes",
        title="Prediabetes",
        organization="Centers for Disease Control and Prevention",
        year=2024,
        source_type="public_health_page",
        url="https://www.cdc.gov/diabetes/basics/prediabetes.html",
        content=(
            "Prediabetes is a warning sign that blood sugar is higher than normal "
            "but not yet in the diabetes range. Lifestyle changes such as physical "
            "activity and weight management can reduce risk of progression."
        ),
    ),
    TrustedSourceDocument(
        citation_id="who_physical_activity_2020",
        title="WHO Guidelines on Physical Activity and Sedentary Behaviour",
        organization="World Health Organization",
        year=2020,
        source_type="guideline",
        url="https://www.who.int/publications/i/item/9789240015128",
        content=(
            "Regular physical activity supports cardiovascular and metabolic health. "
            "Guidelines commonly recommend at least 150 minutes per week of moderate "
            "aerobic activity for adults, with additional muscle strengthening work."
        ),
    ),
    TrustedSourceDocument(
        citation_id="nih_lipid_panel",
        title="Lipid Panel",
        organization="MedlinePlus / NIH",
        year=2024,
        source_type="patient_education",
        url="https://medlineplus.gov/lab-tests/lipid-panel/",
        content=(
            "A lipid panel measures cholesterol and triglycerides. The results help "
            "clinicians evaluate cardiovascular risk and decide whether lifestyle "
            "change or medication may be appropriate."
        ),
    ),
    TrustedSourceDocument(
        citation_id="niddk_kidney_tests",
        title="Kidney Function Tests",
        organization="NIDDK / NIH",
        year=2024,
        source_type="patient_education",
        url="https://www.niddk.nih.gov/health-information/diagnostic-tests/kidney-function-tests",
        content=(
            "Creatinine and blood urea nitrogen are commonly used to assess kidney "
            "function. Interpretation should consider sex, age, muscle mass, and the "
            "laboratory reference range."
        ),
    ),
    TrustedSourceDocument(
        citation_id="nih_cbc",
        title="Complete Blood Count (CBC)",
        organization="MedlinePlus / NIH",
        year=2024,
        source_type="patient_education",
        url="https://medlineplus.gov/lab-tests/complete-blood-count-cbc/",
        content=(
            "A CBC includes hemoglobin, white blood cells, platelets, and related "
            "indices. Interpretation depends on age, sex, pregnancy status, and the "
            "laboratory's own reference intervals."
        ),
    ),
]


TEST_ALIASES: dict[str, list[str]] = {
    "ldl_cholesterol": [r"ldl(?:\s+cholesterol)?", r"low[-\s]?density lipoprotein"],
    "hdl_cholesterol": [r"hdl(?:\s+cholesterol)?", r"high[-\s]?density lipoprotein"],
    "triglycerides": [r"triglycerides?", r"tg"],
    "total_cholesterol": [r"total\s+cholesterol", r"cholesterol"],
    "fasting_glucose": [r"fasting\s+glucose", r"glucose", r"blood\s+glucose", r"plasma\s+glucose"],
    "hba1c": [r"hba1c", r"glycated\s+hemoglobin", r"hemoglobin\s+a1c"],
    "hemoglobin": [r"hemoglobin", r"\bhb\b"],
    "wbc": [r"wbc", r"white\s+blood\s+cells?", r"leukocytes?", r"total\s+leukocyte\s+count"],
    "rbc": [r"rbc", r"red\s+blood\s+cells?", r"total\s+rbc\s+count"],
    "platelets": [r"platelets?", r"platelet\s+count"],
    "neutrophils": [r"neutrophils?"],
    "lymphocytes": [r"lymphocytes?"],
    "eosinophils": [r"eosinophils?"],
    "monocytes": [r"monocytes?"],
    "basophils": [r"basophils?"],
    "hematocrit": [r"hematocrit\s+value", r"\bhct\b"],
    "mcv": [r"mean\s+corpuscular\s+volume", r"\bmcv\b"],
    "mchc": [r"mean\s+cell\s+haemoglobin\s+con", r"\bmchc\b"],
    "mch": [r"mean\s+cell\s+haemoglobin", r"\bmch\b"],
    "creatinine": [r"creatinine"],
    "bun": [r"\bbun\b", r"blood\s+urea\s+nitrogen"],
    "ast": [r"\bast\b", r"sgot"],
    "alt": [r"\balt\b", r"sgpt"],
    "sodium": [r"sodium", r"na\+"],
    "potassium": [r"potassium", r"k\+"],
    "tsh": [r"\btsh\b", r"thyroid\s+stimulating\s+hormone"],
}


def _safe_float(value: str | float | int | None) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    try:
        return float(str(value).replace(",", "").strip())
    except ValueError:
        return None


def _normalize_label(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


def _extract_first(patterns: Iterable[str], text: str) -> Optional[str]:
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0)
    return None


class MedicalTextExtractor:
    """Extract text from PDF, DOCX and TXT reports."""

    @staticmethod
    def extract(file_path: str) -> tuple[str, str]:
        suffix = Path(file_path).suffix.lower()
        if suffix == ".pdf":
            return MedicalTextExtractor.extract_pdf(file_path)
        if suffix == ".docx":
            return MedicalTextExtractor.extract_docx(file_path)
        if suffix == ".txt":
            return MedicalTextExtractor.extract_txt(file_path)
        raise ValueError("Unsupported file format.")

    @staticmethod
    def extract_pdf(file_path: str) -> tuple[str, str]:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"

        text = text.strip()
        if len(text) >= 120:
            return text, "pdf_text"

        if fitz is None:
            return text, "pdf_text"

        ocr_text = []
        document = fitz.open(file_path)
        try:
            for page in document:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_bytes = pix.tobytes("png")
                from PIL import Image
                import io

                with Image.open(io.BytesIO(image_bytes)) as image:
                    ocr_text.append(pytesseract.image_to_string(image))
        finally:
            document.close()

        combined = "\n".join(part.strip() for part in ocr_text if part.strip()).strip()
        return (combined or text), "pdf_ocr"

    @staticmethod
    def extract_docx(file_path: str) -> tuple[str, str]:
        if docx is not None:
            document = docx.Document(file_path)  # type: ignore[union-attr]
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            return "\n".join(parts).strip(), "docx"

        with zipfile.ZipFile(file_path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        text = re.sub(r"<[^>]+>", " ", xml)
        text = re.sub(r"\s+", " ", text).strip()
        return text, "docx_xml"

    @staticmethod
    def extract_txt(file_path: str) -> tuple[str, str]:
        return Path(file_path).read_text(encoding="utf-8", errors="ignore").strip(), "txt"


class ReferenceRangeService:
    """Determine reference ranges and abnormal status for common lab tests."""

    @staticmethod
    def parse_reference_range(reference_range: Optional[str]) -> dict[str, Optional[float | str]]:
        if not reference_range:
            return {"min": None, "max": None, "text": None}

        cleaned = reference_range.strip()
        cleaned = re.sub(r"(?<=\d),(?=\d)", "", cleaned)
        range_match = re.search(
            r"(?P<prefix><=|>=|<|>)?\s*(?P<first>\d+(?:\.\d+)?)"
            r"(?:\s*[-]\s*(?P<second>\d+(?:\.\d+)?))?",
            cleaned,
        )
        if not range_match:
            return {"min": None, "max": None, "text": cleaned}

        prefix = range_match.group("prefix")
        first = _safe_float(range_match.group("first"))
        second = _safe_float(range_match.group("second"))

        if prefix in {"<"}:
            return {"min": None, "max": first, "text": cleaned}
        if prefix in {">"}:
            return {"min": first, "max": None, "text": cleaned}
        if second is not None:
            return {"min": first, "max": second, "text": cleaned}

        return {"min": first, "max": first, "text": cleaned}

    @staticmethod
    def infer(test_name: str, value: float, unit: Optional[str], sex: Optional[str]) -> tuple[str, str, str]:
        key = _normalize_label(test_name)
        sex_key = (sex or "").strip().lower()

        def clip_range(text: str, status: str, explanation: str) -> tuple[str, str, str]:
            return text, status, explanation

        if "ldl" in key:
            reference = "<100 mg/dL"
            if value >= 190:
                return clip_range(reference, "CRITICAL", "Markedly elevated LDL cholesterol.")
            if value >= 160:
                return clip_range(reference, "HIGH", "High LDL cholesterol.")
            if value >= 130:
                return clip_range(reference, "HIGH", "Above the optimal LDL target.")
            if value >= 100:
                return clip_range(reference, "BORDERLINE", "Above optimal but not severely elevated.")
            return clip_range(reference, "NORMAL", "LDL is within the optimal range.")

        if "hdl" in key:
            reference = "Men: >=40 mg/dL; Women: >=50 mg/dL"
            threshold = 50 if "female" in sex_key or "woman" in sex_key else 40
            if value < threshold:
                return clip_range(reference, "LOW", "HDL is below the preferred threshold.")
            if value >= 60:
                return clip_range(reference, "NORMAL", "HDL is favorable.")
            return clip_range(reference, "NORMAL", "HDL is acceptable.")

        if "triglyceride" in key:
            reference = "<150 mg/dL"
            if value >= 500:
                return clip_range(reference, "CRITICAL", "Very high triglycerides.")
            if value >= 200:
                return clip_range(reference, "HIGH", "High triglycerides.")
            if value >= 150:
                return clip_range(reference, "BORDERLINE", "Borderline high triglycerides.")
            return clip_range(reference, "NORMAL", "Triglycerides are within range.")

        if "fasting glucose" in key or key in {"glucose", "blood glucose", "plasma glucose"}:
            reference = "70-99 mg/dL fasting"
            if value >= 200:
                return clip_range(reference, "CRITICAL", "Severely elevated glucose.")
            if value >= 126:
                return clip_range(reference, "HIGH", "Fasting glucose is in the diabetic range.")
            if value >= 100:
                return clip_range(reference, "BORDERLINE", "Fasting glucose is in the impaired range.")
            return clip_range(reference, "NORMAL", "Glucose is within the expected fasting range.")

        if "hba1c" in key:
            reference = "<5.7%"
            if value >= 6.5:
                return clip_range(reference, "HIGH", "HbA1c is in the diabetic range.")
            if value >= 5.7:
                return clip_range(reference, "BORDERLINE", "HbA1c is in the prediabetes range.")
            return clip_range(reference, "NORMAL", "HbA1c is within the normal range.")

        if "hemoglobin" in key or key == "hb":
            if "female" in sex_key or "woman" in sex_key:
                reference = "12.0-15.5 g/dL"
                if value < 12.0:
                    return clip_range(reference, "LOW", "Hemoglobin is low for a female reference interval.")
                if value > 15.5:
                    return clip_range(reference, "HIGH", "Hemoglobin is above the female reference interval.")
                return clip_range(reference, "NORMAL", "Hemoglobin is within the female reference interval.")
            reference = "13.5-17.5 g/dL"
            if value < 13.5:
                return clip_range(reference, "LOW", "Hemoglobin is low for a male reference interval.")
            if value > 17.5:
                return clip_range(reference, "HIGH", "Hemoglobin is above the male reference interval.")
            return clip_range(reference, "NORMAL", "Hemoglobin is within the male reference interval.")

        if "platelet" in key:
            reference = "150-450 x10^3/uL"
            if value < 150:
                return clip_range(reference, "LOW", "Platelet count is low.")
            if value > 450:
                return clip_range(reference, "HIGH", "Platelet count is elevated.")
            return clip_range(reference, "NORMAL", "Platelet count is within range.")

        if "wbc" in key or "white blood" in key:
            reference = "4.0-11.0 x10^3/uL"
            if value < 4.0:
                return clip_range(reference, "LOW", "White blood cell count is low.")
            if value > 11.0:
                return clip_range(reference, "HIGH", "White blood cell count is elevated.")
            return clip_range(reference, "NORMAL", "White blood cell count is within range.")

        if "creatinine" in key:
            reference = "0.6-1.3 mg/dL"
            if value < 0.6:
                return clip_range(reference, "LOW", "Creatinine is below the usual range.")
            if value > 1.3:
                return clip_range(reference, "HIGH", "Creatinine is above the usual range.")
            return clip_range(reference, "NORMAL", "Creatinine is within the usual range.")

        if "bun" in key:
            reference = "7-20 mg/dL"
            if value < 7:
                return clip_range(reference, "LOW", "BUN is low.")
            if value > 20:
                return clip_range(reference, "HIGH", "BUN is elevated.")
            return clip_range(reference, "NORMAL", "BUN is within range.")

        if "ast" in key or "alt" in key:
            reference = "Within laboratory reference range"
            if value > 2 * 40:
                return clip_range(reference, "HIGH", "Transaminase elevation is significant.")
            if value > 40:
                return clip_range(reference, "BORDERLINE", "Transaminase is mildly elevated.")
            return clip_range(reference, "NORMAL", "Transaminase is within range.")

        if "sodium" in key:
            reference = "135-145 mEq/L"
            if value < 135:
                return clip_range(reference, "LOW", "Sodium is low.")
            if value > 145:
                return clip_range(reference, "HIGH", "Sodium is high.")
            return clip_range(reference, "NORMAL", "Sodium is within range.")

        if "potassium" in key:
            reference = "3.5-5.0 mEq/L"
            if value < 3.5:
                return clip_range(reference, "LOW", "Potassium is low.")
            if value > 5.5:
                return clip_range(reference, "CRITICAL", "Potassium is severely elevated.")
            if value > 5.0:
                return clip_range(reference, "HIGH", "Potassium is elevated.")
            return clip_range(reference, "NORMAL", "Potassium is within range.")

        if "tsh" in key:
            reference = "0.4-4.0 uIU/mL"
            if value < 0.4:
                return clip_range(reference, "LOW", "TSH is low.")
            if value > 4.0:
                return clip_range(reference, "HIGH", "TSH is elevated.")
            return clip_range(reference, "NORMAL", "TSH is within range.")

        return "Unknown", "UNKNOWN", "No trusted reference range available for this test."


class MedicalReportParser:
    """Parse extracted text into report facts and findings."""

    @staticmethod
    def extract_patient_info(text: str, report_type: Optional[str] = None) -> PatientInfo:
        lines = text.splitlines()
        name = None
        age_text = None
        sex = None
        patient_id = None
        report_date = None
        referring_physician = None

        for i, raw_line in enumerate(lines):
            line = raw_line.strip()
            lower = line.lower()
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""

            if not name and ("patient name" in lower or (lower.endswith("name:") and "patient" in lower)):
                candidate = next_line if line.endswith(":") and not next_line.lower().startswith("age") else line.split(":", 1)[-1].strip()
                if candidate:
                    name = candidate.split("\n")[0].strip()

            if not age_text and ("age" in lower and ("sex" in lower or "/" in lower)):
                candidate = next_line if line.endswith(":") else line.split(":", 1)[-1].strip()
                if candidate:
                    age_text = candidate.split("\n")[0].strip()

            if not sex and ((lower.startswith("sex") and ":" in line) or ("age / sex" in lower)):
                candidate = next_line if line.endswith(":") else line.split(":", 1)[-1].strip()
                if candidate:
                    sex = candidate.split("\n")[0].strip()

            if not patient_id and ("reg" in lower and "no" in lower):
                candidate = next_line if line.endswith(":") else line.split(":", 1)[-1].strip()
                if candidate:
                    patient_id = candidate.split("\n")[0].strip()

            if not report_date and lower.startswith("reported on") and ":" in line:
                candidate = next_line if line.endswith(":") else line.split(":", 1)[-1].strip()
                if candidate:
                    report_date = candidate.split("\n")[0].strip()

            if not report_date and lower.startswith("collected on") and ":" in line:
                candidate = next_line if line.endswith(":") else line.split(":", 1)[-1].strip()
                if candidate and not report_date:
                    report_date = candidate.split("\n")[0].strip()

            if not report_date and lower.startswith("registered on") and ":" in line:
                candidate = next_line if line.endswith(":") else line.split(":", 1)[-1].strip()
                if candidate and not report_date:
                    report_date = candidate.split("\n")[0].strip()

            if not referring_physician and ("referred by" in lower or "referring" in lower):
                candidate = next_line if line.endswith(":") else line.split(":", 1)[-1].strip()
                if candidate:
                    referring_physician = candidate.split("\n")[0].strip()

        age = None
        if age_text:
            age_match = re.search(r"\d{1,3}", age_text)
            if age_match:
                age = int(age_match.group(0))

        if sex and ("yrs" in sex.lower() or "yr" in sex.lower() or "/" in sex.lower()):
            parts = re.split(r"\s*/\s*", sex)
            for part in parts:
                part = part.strip()
                if re.match(r"^(male|female|m|f)$", part, re.IGNORECASE):
                    sex = part.upper() if part.lower() in ("m", "f") else part
                    break
            else:
                if parts:
                    sex = parts[-1].strip()

        return PatientInfo(
            name=name,
            age=age,
            sex=sex,
            patient_id=patient_id,
            report_date=report_date,
            referring_physician=referring_physician,
            report_type=report_type,
        )

    @staticmethod
    def _match_test_name(line: str) -> Optional[str]:
        normalized = _normalize_label(line)
        for canonical_name, patterns in TEST_ALIASES.items():
            if any(re.search(pattern, normalized, re.IGNORECASE) for pattern in patterns):
                return canonical_name
        return None

    @staticmethod
    def _extract_reference_from_line(line: str) -> Optional[str]:
        match = re.search(
            r"(?:reference|normal)\s*(?:range)?\s*[:=]?\s*([<>=]?\s*\d{1,3}(?:,\d{3})*(?:\.\d+)?(?:\s*[-–]\s*\d{1,3}(?:,\d{3})*(?:\.\d+)?)?\s*[A-Za-z/%^\d\-\.\s,]*)",
            line,
            re.IGNORECASE,
        )
        if match:
            return match.group(1).strip()

        match = re.search(
            r"\(\s*([<>=]?\s*\d{1,3}(?:,\d{3})*(?:\.\d+)?(?:\s*[-–]\s*\d{1,3}(?:,\d{3})*(?:\.\d+)?)?\s*[A-Za-z/%^\d\-\.\s,]*)\s*\)",
            line,
        )
        if match:
            return match.group(1).strip()

        return None

    @staticmethod
    def extract_findings(text: str, patient_info: PatientInfo) -> list[ReportFinding]:
        findings: list[ReportFinding] = []
        seen: set[tuple[str, str]] = set()
        lines = text.splitlines()

        last_test_name: Optional[str] = None
        last_value_line_idx: Optional[int] = None

        def flush_test(
            test_name: Optional[str],
            value_line_idx: Optional[int],
            value: float,
            unit: Optional[str],
            reference_range: Optional[str],
        ) -> None:
            if test_name is None or value_line_idx is None:
                return
            canonical = MedicalReportParser._match_test_name(test_name)
            if canonical is None:
                return

            inferred_reference, status, interpretation = ReferenceRangeService.infer(
                canonical.replace("_", " "),
                value,
                unit,
                patient_info.sex,
            )

            if reference_range:
                parsed = ReferenceRangeService.parse_reference_range(reference_range)
                if parsed["max"] is not None and value > float(parsed["max"]):
                    status = "HIGH"
                    interpretation = "Above the reference range reported by the laboratory."
                elif parsed["min"] is not None and value < float(parsed["min"]):
                    status = "LOW"
                    interpretation = "Below the reference range reported by the laboratory."
                else:
                    status = "NORMAL"
                    interpretation = "Within the reference range reported by the laboratory."

            key = (canonical, f"{value}")
            if key in seen:
                return
            seen.add(key)

            findings.append(
                ReportFinding(
                    test_name=canonical.replace("_", " ").title(),
                    value=round(value, 2) if value.is_integer() is False else int(value),
                    unit=unit,
                    reference_range=reference_range or inferred_reference,
                    status=status,
                    interpretation=interpretation,
                    evidence_ids=[],
                )
            )

        i = 0
        while i < len(lines):
            raw_line = lines[i]
            line = raw_line.strip()
            if not line or len(line) < 3:
                i += 1
                continue

            canonical = MedicalReportParser._match_test_name(line)
            if canonical is not None:
                last_test_name = line
                last_value_line_idx = i + 1 if (i + 1 < len(lines) and lines[i + 1].strip()) else i
                i += 1
                continue

            value_match = re.search(r"(-?\d{1,3}(?:,\d{3})*(?:\.\d+)?)", line)
            if value_match and last_test_name is not None:
                raw_value = value_match.group(1)
                value = _safe_float(raw_value.replace(",", ""))
                if value is not None:
                    unit_match = re.search(
                        r"\b(mg/dL|g/dL|mmol/L|x10\^?3/?uL|K/?uL|cumm|lakhs/cumm|million/cumm|fL|Pg|mEq/L|IU/mL|uIU/mL|%)\b",
                        line,
                        re.IGNORECASE,
                    )
                    unit = unit_match.group(1) if unit_match else None

                    reference_range = None
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if next_line.lower().startswith("reference") or next_line.lower().startswith("normal"):
                            reference_range = MedicalReportParser._extract_reference_from_line(next_line)
                        elif re.match(r"^[<>=]?\s*\d", next_line):
                            reference_range = MedicalReportParser._extract_reference_from_line(next_line)

                    if not reference_range:
                        reference_range = MedicalReportParser._extract_reference_from_line(line)

                    flush_test(last_test_name, last_value_line_idx, value, unit, reference_range)
                    last_test_name = None
                    last_value_line_idx = None
                    i += 1
                    continue

            if last_test_name is not None and i > last_value_line_idx:
                lower = line.lower()
                if lower.startswith("reference") or lower.startswith("normal") or re.match(r"^[<>=]?\s*\d", lower):
                    reference_range = MedicalReportParser._extract_reference_from_line(line)
                    if reference_range:
                        last_test_name = None
                        last_value_line_idx = None

            i += 1

        return findings

    @staticmethod
    def detect_conditions(findings: list[ReportFinding]) -> list[DetectedCondition]:
        by_name = {finding.test_name.lower(): finding for finding in findings}
        conditions: list[DetectedCondition] = []

        ldl = by_name.get("ldl cholesterol")
        hdl = by_name.get("hdl cholesterol")
        triglycerides = by_name.get("triglycerides")
        glucose = by_name.get("fasting glucose")
        hba1c = by_name.get("hba1c")
        hemoglobin = by_name.get("hemoglobin")

        if ldl and ldl.status in {"HIGH", "BORDERLINE", "CRITICAL"}:
            conditions.append(DetectedCondition(name="Possible dyslipidemia", confidence=0.86))
        if hdl and hdl.status == "LOW":
            conditions.append(DetectedCondition(name="Possible atherogenic lipid pattern", confidence=0.78))
        if triglycerides and triglycerides.status in {"HIGH", "BORDERLINE", "CRITICAL"}:
            conditions.append(DetectedCondition(name="Possible hypertriglyceridemia", confidence=0.74))
        if glucose and glucose.status in {"HIGH", "BORDERLINE", "CRITICAL"}:
            conditions.append(DetectedCondition(name="Possible impaired fasting glucose", confidence=0.71))
        if hba1c and hba1c.status in {"BORDERLINE", "HIGH", "CRITICAL"}:
            conditions.append(DetectedCondition(name="Possible dysglycemia", confidence=0.69))
        if hemoglobin and hemoglobin.status == "LOW":
            conditions.append(DetectedCondition(name="Possible anemia", confidence=0.75))

        if len(conditions) >= 2 and any(
            condition.name in {"Possible dyslipidemia", "Possible impaired fasting glucose"}
            for condition in conditions
        ):
            conditions.append(DetectedCondition(name="Possible metabolic risk pattern", confidence=0.58))

        return conditions

    @staticmethod
    def build_important_terms(findings: list[ReportFinding]) -> list[MedicalTerm]:
        terms: list[MedicalTerm] = []
        seen: set[str] = set()

        def add(term: str, meaning: str) -> None:
            if term.lower() in seen:
                return
            seen.add(term.lower())
            terms.append(MedicalTerm(term=term, meaning=meaning))

        test_names = {finding.test_name.lower() for finding in findings}
        if "ldl cholesterol" in test_names:
            add("LDL", "Low-density lipoprotein cholesterol. Higher values are associated with cardiovascular risk.")
        if "hdl cholesterol" in test_names:
            add("HDL", "High-density lipoprotein cholesterol. Lower values are generally less favorable.")
        if "triglycerides" in test_names:
            add("Triglycerides", "A blood fat measured in lipid panels; elevated values can increase cardiometabolic risk.")
        if "fasting glucose" in test_names:
            add("Fasting glucose", "Blood glucose measured after fasting; used to screen for abnormal glucose regulation.")
        if "hba1c" in test_names:
            add("HbA1c", "An estimate of average blood glucose over approximately the prior 2 to 3 months.")
        if "hemoglobin" in test_names:
            add("Hemoglobin", "A red blood cell protein that carries oxygen; low values can indicate anemia.")

        add("Reference range", "The interval a laboratory uses to help interpret whether a result is low, normal, or high.")
        add("Clinical correlation", "Interpretation should be combined with symptoms, history, and the laboratory context.")
        return terms


class MedicalEmbeddingStore:
    """Embedding and vector-store wrapper for trusted source retrieval."""

    def __init__(self) -> None:
        self.embedding_manager = EmbeddingsManager(model_name=settings.MEDICAL_EMBEDDING_MODEL, use_gpu=False)
        self.vector_store = VectorStoreManager(
            index_dir=settings.VECTORSTORE_DIR,
            embedding_dim=self.embedding_manager.get_embedding_dimension(),
        )
        self.retriever = Retriever(self.vector_store, self.embedding_manager, k=4)
        self._loaded = False

    def _source_documents(self) -> list[dict[str, Any]]:
        docs: list[dict[str, Any]] = []
        for source in TRUSTED_SOURCE_LIBRARY:
            docs.append(
                {
                    "content": source.content,
                    "metadata": {
                        "citation_id": source.citation_id,
                        "source": source.organization,
                        "title": source.title,
                        "year": source.year,
                        "source_type": source.source_type,
                        "url": source.url,
                    },
                }
            )
        return docs

    async def ensure_loaded(self) -> None:
        if self._loaded and self.vector_store.get_document_count() > 0:
            return

        index_name = settings.MEDICAL_VECTORSTORE_NAME
        loaded = await self.vector_store.load_index(index_name)
        if loaded and self.vector_store.get_document_count() > 0:
            self._loaded = True
            return

        from langchain_core.documents import Document

        documents = [
            Document(page_content=doc["content"], metadata=doc["metadata"])
            for doc in self._source_documents()
        ]
        texts, embeddings = await self.embedding_manager.embed_documents_batch(documents)
        _ = texts
        self.vector_store.add_documents(documents, embeddings)
        self.vector_store.save_index(index_name)
        self._loaded = True

    @staticmethod
    def generate_queries(findings: list[ReportFinding]) -> list[str]:
        queries: list[str] = []

        for finding in findings:
            name = finding.test_name.lower()
            if "ldl" in name:
                queries.extend(["LDL cholesterol clinical guideline", "LDL cholesterol interpretation"])
            elif "hdl" in name:
                queries.extend(["HDL cholesterol guideline", "HDL cholesterol low interpretation"])
            elif "triglycerides" in name:
                queries.extend(["triglyceride clinical interpretation", "triglyceride guideline"])
            elif "fasting glucose" in name or "hba1c" in name:
                queries.extend(["prediabetes guideline fasting glucose", "HbA1c interpretation guideline"])
            elif "hemoglobin" in name:
                queries.extend(["hemoglobin normal range guideline", "anemia evaluation guideline"])
            elif "creatinine" in name or "bun" in name:
                queries.extend(["kidney function tests interpretation guideline", "creatinine reference range guideline"])
            else:
                queries.extend([f"{finding.test_name} guideline", f"{finding.test_name} interpretation"])

        if not queries:
            queries = ["medical report interpretation guideline"]

        deduped: list[str] = []
        seen: set[str] = set()
        for query in queries:
            key = query.lower()
            if key in seen:
                continue
            seen.add(key)
            deduped.append(query)

        return deduped

    async def retrieve(self, queries: list[str]) -> tuple[list[EvidenceSource], list[dict[str, Any]]]:
        await self.ensure_loaded()

        evidence_by_id: dict[str, EvidenceSource] = {}
        raw_hits: list[dict[str, Any]] = []
        best_scores: dict[str, float] = {}
        score_sums: dict[str, float] = {}
        query_counts: dict[str, int] = {}

        for query in queries:
            results = await self.retriever.retrieve(query, k=3, similarity_threshold=0.25)
            for document, score in results:
                metadata = document.metadata or {}
                citation_id = str(metadata.get("citation_id") or metadata.get("source") or "unknown_source")
                source = EvidenceSource(
                    citation_id=citation_id,
                    title=str(metadata.get("title", "Unknown source")),
                    organization=str(metadata.get("source", "Unknown organization")),
                    year=int(metadata.get("year", 0) or 0),
                    source_type=str(metadata.get("source_type", "document")),
                    url=str(metadata.get("url", "")),
                    chunk_id=str(metadata.get("chunk_id") or f"{citation_id}_chunk"),
                    excerpt=document.page_content[:800].strip(),
                )
                evidence_by_id[source.citation_id] = source
                raw_hits.append(
                    {
                        "citation_id": source.citation_id,
                        "score": score,
                        "query": query,
                        "metadata": metadata,
                        "excerpt": source.excerpt,
                    }
                )
                if citation_id not in best_scores or score > best_scores[citation_id]:
                    best_scores[citation_id] = score
                score_sums[citation_id] = score_sums.get(citation_id, 0.0) + score
                query_counts[citation_id] = query_counts.get(citation_id, 0) + 1

        total_queries = max(len(queries), 1)
        ranked = []
        for citation_id, best_score in best_scores.items():
            avg_score = score_sums.get(citation_id, 0.0) / query_counts.get(citation_id, 1)
            query_fraction = query_counts.get(citation_id, 0) / total_queries
            combined_score = best_score * 0.6 + avg_score * 0.3 + query_fraction * 0.1
            if combined_score >= 0.45:
                ranked.append((combined_score, citation_id))

        ranked.sort(key=lambda item: item[0], reverse=True)
        ranked_ids = [citation_id for _, citation_id in ranked[:5]]

        ordered_sources = [
            evidence_by_id[citation_id]
            for citation_id in ranked_ids
            if citation_id in evidence_by_id
        ]
        return ordered_sources, raw_hits

    @staticmethod
    def format_context(sources: list[EvidenceSource]) -> str:
        parts: list[str] = []
        for source in sources:
            parts.append(
                "\n".join(
                    [
                        f"SOURCE_ID: {source.citation_id}",
                        f"TITLE: {source.title}",
                        f"ORGANIZATION: {source.organization}",
                        f"YEAR: {source.year}",
                        f"TYPE: {source.source_type}",
                        f"URL: {source.url}",
                        f"EXCERPT: {source.excerpt}",
                    ]
                )
            )
        return "\n\n".join(parts)


class GeminiService:
    """Backend-only Gemini integration."""

    def __init__(self) -> None:
        self.api_key = settings.GEMINI_API_KEY.strip()
        self.model = settings.GEMINI_MODEL.strip() or "gemini-3.6-flash"
        self.enabled = bool(self.api_key and genai is not None and types is not None)
        self.client = genai.Client(api_key=self.api_key) if self.enabled else None

    def configuration_state(self) -> dict[str, str]:
        return {
            "status": "llm_not_configured",
            "message": (
                "Gemini API key is not configured. Report extraction and RAG retrieval are "
                "available, but final LLM generation requires GEMINI_API_KEY."
            ),
        }

    def system_prompt(self) -> str:
        return (
            "You are MedIntel AI, a medical report analysis assistant.\n\n"
            "Analyze the complete medical report provided below.\n\n"
            "Your job is to understand the report accurately and return a structured analysis.\n\n"
            "IMPORTANT RULES:\n"
            "1. Use ONLY information actually present in the uploaded report.\n"
            "2. Do not invent patient information.\n"
            "3. Do not invent laboratory values.\n"
            "4. Do not invent reference ranges.\n"
            "5. Do not invent symptoms, medical history, medications or diagnoses.\n"
            "6. Identify all clinically relevant abnormal findings present in the report.\n"
            "7. Also mention important normal findings when they provide useful context.\n"
            "8. If a laboratory value has a reference range in the report, compare the value against that reference range.\n"
            "9. Preserve the original units.\n"
            "10. Clearly distinguish an abnormal laboratory finding from a confirmed medical diagnosis.\n"
            "11. Do not make a definitive diagnosis based only on laboratory results.\n"
            "12. Explain possible clinical significance cautiously.\n"
            "13. Indicate when clinical correlation with a healthcare professional is required.\n"
            "14. Generate practical and medically appropriate follow-up recommendations.\n"
            "15. Do not omit important findings just because there are many of them.\n"
            "16. Analyze the COMPLETE report rather than focusing on only one value.\n"
            "17. Return the response in the required JSON structure.\n\n"
            "The uploaded report is the source of truth for patient-specific facts.\n"
            "Do not merge external knowledge with the report or invent information between them."
        )

    def build_report_prompt(self, report_text: str, report_type: str) -> str:
        return (
            f"Report Type: {report_type or 'Medical Report'}\n\n"
            "Complete Report Content:\n"
            f"{report_text}\n\n"
            "Return JSON matching this structure:\n"
            "{\n"
            '  "patient_information": {\n'
            '    "name": "",\n'
            '    "age": "",\n'
            '    "sex": "",\n'
            '    "patient_id": "",\n'
            '    "referred_by": "",\n'
            '    "report_date": "",\n'
            '    "report_type": ""\n'
            "  },\n"
            '  "medical_summary": "",\n'
            '  "key_findings": [\n'
            "    {\n"
            '      "parameter": "",\n'
            '      "value": "",\n'
            '      "unit": "",\n'
            '      "reference_range": "",\n'
            '      "status": "HIGH | LOW | NORMAL | ABNORMAL | NOT_AVAILABLE",\n'
            '      "interpretation": ""\n'
            "    }\n"
            "  ],\n"
            '  "detected_conditions": [\n'
            "    {\n"
            '      "name": "",\n'
            '      "description": "",\n'
            '      "clinical_significance": "",\n'
            '      "clinical_correlation_required": true\n'
            "    }\n"
            "  ],\n"
            '  "recommendations": [\n'
            '    ""\n'
            "  ],\n"
            '  "important_terms": [\n'
            "    {\n"
            '      "term": "",\n'
            '      "explanation": ""\n'
            "    }\n"
            "  ]\n"
            "}\n"
            "If a section cannot be determined from the report, return an empty array or empty string rather than inventing information."
        )

    async def analyze_report_text(
        self,
        report_text: str,
        report_type: str,
    ) -> dict[str, Any]:
        if not self.enabled or self.client is None:
            return self.configuration_state()

        prompt = self.build_report_prompt(report_text, report_type)
        try:
            response = await asyncio.to_thread(
                self.client.models.generate_content,
                model=self.model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    system_instruction=self.system_prompt(),
                    response_mime_type="application/json",
                    temperature=0.2,
                ),
            )
        except Exception as exc:
            logger.warning("Gemini report analysis failed: %s", exc)
            return {
                "status": "llm_error",
                "message": "AI analysis could not be completed. Please try again.",
            }

        response_text = getattr(response, "text", "") or ""
        if not response_text.strip():
            return {
                "status": "llm_error",
                "message": "AI analysis could not be completed. Please try again.",
            }

        try:
            parsed = GeminiReportAnalysis.model_validate_json(response_text)
        except Exception as exc:
            logger.warning("Gemini report analysis validation failed: %s", exc)
            return {
                "status": "llm_error",
                "message": "AI analysis could not be completed. Please try again.",
            }

        return parsed.model_dump()

    def build_prompt(self, report_facts: dict[str, Any], evidence_context: str) -> str:
        return (
            "REPORT FACTS:\n"
            f"{json.dumps(report_facts, indent=2, ensure_ascii=True)}\n\n"
            "RETRIEVED EVIDENCE:\n"
            f"{evidence_context}\n\n"
            "Return JSON matching this structure:\n"
            "{\n"
            '  "summary": "...",\n'
            '  "findings": [\n'
            "    {\n"
            '      "test_name": "LDL Cholesterol",\n'
            '      "value": 142,\n'
            '      "unit": "mg/dL",\n'
            '      "reference_range": "<100",\n'
            '      "status": "HIGH",\n'
            '      "interpretation": "...",\n'
            '      "evidence_ids": ["source_1"]\n'
            "    }\n"
            "  ],\n"
            '  "possible_conditions": [\n'
            "    {\n"
            '      "name": "...",\n'
            '      "explanation": "...",\n'
            '      "evidence_ids": ["source_1"],\n'
            '      "clinical_correlation_required": true\n'
            "    }\n"
            "  ],\n"
            '  "recommendations": [\n'
            "    {\n"
            '      "text": "...",\n'
            '      "evidence_ids": ["source_1"]\n'
            "    }\n"
            "  ],\n"
            '  "important_terms": [\n'
            "    {\n"
            '      "term": "...",\n'
            '      "definition": "...",\n'
            '      "evidence_ids": ["source_1"]\n'
            "    }\n"
            "  ],\n"
            '  "limitations": ["..."]\n'
            "}\n"
            "Only reference evidence IDs that appear in the retrieved evidence block."
        )

    async def generate(
        self,
        report_facts: dict[str, Any],
        evidence_sources: list[EvidenceSource],
    ) -> dict[str, Any]:
        if not self.enabled or self.client is None:
            return self.configuration_state()

        evidence_context = MedicalEmbeddingStore.format_context(evidence_sources)
        prompt = self.build_prompt(report_facts, evidence_context)

        schema = GeminiAnalysisResponse
        try:
            response = await asyncio.to_thread(
                self.client.models.generate_content,
                model=self.model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    system_instruction=self.system_prompt(),
                    response_mime_type="application/json",
                    response_schema=schema,
                    temperature=0.2,
                ),
            )
        except Exception as exc:
            logger.warning("Gemini generation failed: %s", exc)
            return {
                "status": "llm_error",
                "message": "AI analysis could not be completed. Please try again.",
            }

        response_text = getattr(response, "text", "") or ""
        if not response_text.strip():
            return {
                "status": "llm_error",
                "message": "AI analysis could not be completed. Please try again.",
            }

        try:
            parsed = GeminiAnalysisResponse.model_validate_json(response_text)
        except Exception as exc:
            logger.warning("Gemini response validation failed: %s", exc)
            return {
                "status": "llm_error",
                "message": "AI analysis could not be completed. Please try again.",
            }

        return parsed.model_dump()


class GeminiFinding(BaseModel):
    test_name: str
    value: Any
    unit: Optional[str] = None
    reference_range: Optional[str] = None
    status: str
    interpretation: str
    evidence_ids: list[str] = Field(default_factory=list)


class GeminiCondition(BaseModel):
    name: str
    explanation: str
    evidence_ids: list[str] = Field(default_factory=list)
    clinical_correlation_required: bool = True


class GeminiRecommendation(BaseModel):
    text: str
    evidence_ids: list[str] = Field(default_factory=list)


class GeminiTerm(BaseModel):
    term: str
    definition: str
    evidence_ids: list[str] = Field(default_factory=list)


class GeminiAnalysisResponse(BaseModel):
    summary: str
    findings: list[GeminiFinding] = Field(default_factory=list)
    possible_conditions: list[GeminiCondition] = Field(default_factory=list)
    recommendations: list[GeminiRecommendation] = Field(default_factory=list)
    important_terms: list[GeminiTerm] = Field(default_factory=list)
    limitations: list[str] = Field(default_factory=list)


class GeminiPatientInfo(BaseModel):
    name: Optional[str] = None
    age: Optional[str] = None
    sex: Optional[str] = None
    patient_id: Optional[str] = None
    referred_by: Optional[str] = None
    report_date: Optional[str] = None
    report_type: Optional[str] = None


class GeminiKeyFinding(BaseModel):
    parameter: str
    value: str
    unit: Optional[str] = None
    reference_range: Optional[str] = None
    status: str
    interpretation: str


class GeminiDetectedCondition(BaseModel):
    name: str
    description: str
    clinical_significance: str
    clinical_correlation_required: bool = True


class GeminiReportAnalysis(BaseModel):
    patient_information: GeminiPatientInfo
    medical_summary: str
    key_findings: list[GeminiKeyFinding] = Field(default_factory=list)
    detected_conditions: list[GeminiDetectedCondition] = Field(default_factory=list)
    recommendations: list[str] = Field(default_factory=list)
    important_terms: list[dict[str, str]] = Field(default_factory=list)


class MedicalReportAnalysisService:
    """High-level service that orchestrates report extraction, RAG, and Gemini."""

    def __init__(self) -> None:
        self.embedding_store = MedicalEmbeddingStore()
        self.gemini = GeminiService()

    @staticmethod
    def _empty_analysis_state(report: MedicalReport, message: str) -> dict[str, Any]:
        patient_info = PatientInfo(
            name=report.patient_name,
            report_type=report.report_type,
        )
        return ReportAnalysisDetail(
            report_id=report.id,
            status="llm_not_configured",
            message=message,
            llm_status="llm_not_configured",
            patient_info=patient_info,
            summary=message,
            findings=[],
            detected_conditions=[],
            possible_conditions=[],
            recommendations=[],
            important_terms=[],
            evidence_sources=[],
            retrieval_queries=[],
            processing_stage=report.processing_stage,
            educational_use_only=True,
        ).model_dump()

    @staticmethod
    def _serialize_conditions(conditions: list[DetectedCondition]) -> list[dict[str, Any]]:
        return [condition.model_dump() if hasattr(condition, "model_dump") else dict(condition) for condition in conditions]

    @staticmethod
    def _fallback_recommendations(findings: list[ReportFinding]) -> list[str]:
        recommendations: list[str] = []
        names = {finding.test_name.lower(): finding for finding in findings}

        if "ldl cholesterol" in names or "triglycerides" in names or "hdl cholesterol" in names:
            recommendations.append(
                "Discuss lipid results with a clinician and use guideline-based lifestyle measures "
                "for cardiometabolic risk reduction."
            )
        if "fasting glucose" in names or "hba1c" in names:
            recommendations.append(
                "Clinical correlation is important for abnormal glucose results; repeat testing may "
                "be appropriate depending on the overall clinical context."
            )
        if "hemoglobin" in names:
            recommendations.append(
                "Consider evaluation for anemia if the low hemoglobin result is confirmed by the "
                "report and the patient has relevant symptoms or history."
            )
        if not recommendations:
            recommendations.append(
                "Correlate the findings with the full clinical picture and the laboratory reference ranges."
            )
        return recommendations

    async def extract_report(self, db: Session, report: MedicalReport) -> dict[str, Any]:
        report.processing_stage = "PROCESSING"
        report.status = ReportStatus.PROCESSING
        db.commit()
        db.refresh(report)

        text, extraction_method = MedicalTextExtractor.extract(report.file_path)
        if not text.strip():
            report.processing_stage = "FAILED"
            report.status = ReportStatus.FAILED
            report.processing_message = "No readable medical content was found."
            db.commit()
            return {
                "status": "failed",
                "message": "No readable medical content was found.",
            }

        report.extracted_text = text
        report.extracted_data = {
            "extraction_method": extraction_method,
            "text": text,
        }
        report.processing_stage = "EXTRACTED"
        report.summary = report.summary or "Report extracted successfully."
        db.commit()
        db.refresh(report)

        return {
            "status": "completed",
            "extraction_method": extraction_method,
            "text": text,
        }

    async def analyze_report(self, db: Session, report: MedicalReport) -> dict[str, Any]:
        extraction = await self.extract_report(db, report)
        if extraction.get("status") == "failed":
            patient_info = PatientInfo(
                name=report.patient_name,
                report_type=report.report_type,
            )
            analysis = ReportAnalysisDetail(
                report_id=report.id,
                status="failed",
                message=extraction.get("message", "Extraction failed"),
                llm_status="not_run",
                patient_info=patient_info,
                summary=extraction.get("message", "Extraction failed"),
                findings=[],
                detected_conditions=[],
                possible_conditions=[],
                recommendations=[],
                important_terms=[],
                evidence_sources=[],
                retrieval_queries=[],
                processing_stage="FAILED",
                educational_use_only=True,
            ).model_dump()
            report.analysis_payload = analysis
            report.status = ReportStatus.FAILED
            report.processing_stage = "FAILED"
            report.processing_message = extraction.get("message", "Extraction failed")
            db.commit()
            db.refresh(report)
            return analysis

        report_text = extraction.get("text", "")
        report.processing_stage = "ANALYZING_REPORT"
        db.commit()
        db.refresh(report)

        gemini_output = await self.gemini.analyze_report_text(report_text, report.report_type)
        if gemini_output.get("status") == "llm_not_configured":
            analysis = self._empty_analysis_state(
                report,
                gemini_output["message"],
            )
            analysis["processing_stage"] = "COMPLETED"
            report.analysis_payload = analysis
            report.status = ReportStatus.ANALYZED
            report.processing_stage = "COMPLETED"
            report.processing_message = gemini_output["message"]
            db.commit()
            db.refresh(report)
            return analysis

        if gemini_output.get("status") == "llm_error":
            report.processing_message = gemini_output["message"]
            report.status = ReportStatus.FAILED
            report.processing_stage = "FAILED"
            db.commit()

            patient_info = PatientInfo(
                name=report.patient_name,
                report_type=report.report_type,
            )
            analysis = ReportAnalysisDetail(
                report_id=report.id,
                status="failed",
                message=gemini_output["message"],
                llm_status="llm_error",
                patient_info=patient_info,
                summary=gemini_output["message"],
                findings=[],
                detected_conditions=[],
                possible_conditions=[],
                recommendations=[],
                important_terms=[],
                evidence_sources=[],
                retrieval_queries=[],
                processing_stage="FAILED",
                educational_use_only=True,
            ).model_dump()
            report.analysis_payload = analysis
            db.commit()
            db.refresh(report)
            return analysis

        gemini_patient = GeminiPatientInfo.model_validate(gemini_output.get("patient_information", {}))
        patient_info = PatientInfo(
            name=gemini_patient.name,
            age=int(gemini_patient.age) if str(gemini_patient.age or "").isdigit() else None,
            sex=gemini_patient.sex,
            patient_id=gemini_patient.patient_id,
            report_date=gemini_patient.report_date,
            referring_physician=gemini_patient.referred_by,
            report_type=gemini_patient.report_type or report.report_type,
        )

        gemini_findings = [
            GeminiKeyFinding.model_validate(item) for item in gemini_output.get("key_findings", [])
        ]
        findings = [
            ReportFinding(
                test_name=item.parameter,
                value=item.value,
                unit=item.unit,
                reference_range=item.reference_range,
                status=item.status,
                interpretation=item.interpretation,
                evidence_ids=[],
            )
            for item in gemini_findings
        ]

        queries = self._build_queries_from_gemini(gemini_output)
        report.retrieval_queries = queries
        report.processing_stage = "RETRIEVING_EVIDENCE"
        db.commit()
        db.refresh(report)

        evidence_sources, _raw_hits = await self.embedding_store.retrieve(queries)
        report.evidence_sources = [source.model_dump() for source in evidence_sources]
        report.processing_stage = "COMPLETED"
        db.commit()
        db.refresh(report)

        final_findings = []
        for finding in findings:
            final_findings.append(finding.model_dump())

        final_conditions = [
            {
                "name": item.name,
                "description": item.description,
                "clinical_significance": item.clinical_significance,
                "clinical_correlation_required": item.clinical_correlation_required,
                "evidence_ids": [],
            }
            for item in [
                GeminiDetectedCondition.model_validate(item)
                for item in gemini_output.get("detected_conditions", [])
            ]
        ]

        final_recommendations = [
            {"text": item, "evidence_ids": []}
            for item in gemini_output.get("recommendations", [])
        ]

        final_terms = [
            {"term": item.get("term", ""), "meaning": item.get("explanation", ""), "evidence_ids": []}
            for item in gemini_output.get("important_terms", [])
        ]

        analysis = ReportAnalysisDetail(
            report_id=report.id,
            status="completed",
            message=None,
            llm_status="configured",
            patient_info=patient_info,
            summary=gemini_output.get("medical_summary", ""),
            findings=final_findings,
            detected_conditions=[
                DetectedCondition(name=item["name"], confidence=0.0) for item in final_conditions
            ],
            possible_conditions=final_conditions,
            recommendations=final_recommendations,
            important_terms=final_terms,
            evidence_sources=evidence_sources,
            retrieval_queries=queries,
            processing_stage="COMPLETED",
            educational_use_only=True,
        ).model_dump()

        analysis["possible_conditions"] = final_conditions
        analysis["recommendations"] = final_recommendations
        analysis["important_terms"] = final_terms

        report.summary = gemini_output.get("medical_summary", "")
        report.key_findings = final_findings
        report.detected_conditions = final_conditions
        report.recommendations = [item["text"] for item in final_recommendations]
        report.medical_terms = final_terms
        report.analysis_payload = analysis
        report.status = ReportStatus.ANALYZED
        report.processing_stage = "COMPLETED"
        report.processing_message = "Analysis complete."
        db.commit()
        db.refresh(report)

        return analysis

    @staticmethod
    def _build_queries_from_gemini(gemini_output: dict[str, Any]) -> list[str]:
        queries: list[str] = []
        seen: set[str] = set()

        def add_query(text: str) -> None:
            text = text.strip()
            if not text:
                return
            key = text.lower()
            if key in seen:
                return
            seen.add(key)
            queries.append(text)

        def cluster_label(parameter: str) -> str:
            lower = parameter.lower()
            if any(term in lower for term in ["hemoglobin", "rbc", "hematocrit", "hct", "mch", "mchc", "mcv"]):
                return "red blood cell indices"
            if any(term in lower for term in ["wbc", "leukocyte", "neutrophil", "lymphocyte", "eosinophil", "monocyte", "basophil"]):
                return "white blood cell differential"
            if any(term in lower for term in ["platelet"]):
                return "platelet count"
            if any(term in lower for term in ["cholesterol", "ldl", "hdl", "triglyceride", "lipid", "vldl"]):
                return "lipid panel"
            if any(term in lower for term in ["glucose", "hba1c", "blood sugar"]):
                return "glucose metabolism"
            if any(term in lower for term in ["creatinine", "bun", "kidney"]):
                return "kidney function"
            if any(term in lower for term in ["ast", "alt", "liver", "bilirubin"]):
                return "liver function"
            if any(term in lower for term in ["tsh", "thyroid", "t3", "t4"]):
                return "thyroid function"
            return parameter

        clusters: dict[str, list[str]] = {}
        for finding in gemini_output.get("key_findings", []):
            parameter = finding.get("parameter", "")
            if not parameter:
                continue
            cluster = cluster_label(parameter)
            clusters.setdefault(cluster, []).append(parameter)

        for cluster, parameters in clusters.items():
            if len(parameters) == 1:
                add_query(f"{parameters[0]} interpretation")
            else:
                add_query(f"{cluster} interpretation")
                add_query(f"{cluster} guideline")

        for condition in gemini_output.get("detected_conditions", []):
            name = condition.get("name", "")
            if name:
                add_query(name)

        for term in gemini_output.get("important_terms", []):
            term_name = term.get("term", "")
            if term_name:
                add_query(term_name)

        if not queries:
            queries = ["medical report interpretation guideline"]

        return queries[:8]

    @staticmethod
    def _ensure_known_evidence_ids(
        evidence_ids: list[str],
        evidence_sources: list[EvidenceSource],
    ) -> list[str]:
        known = {source.citation_id for source in evidence_sources}
        return [evidence_id for evidence_id in evidence_ids if evidence_id in known]

    async def get_analysis(self, db: Session, report: MedicalReport) -> dict[str, Any]:
        if report.analysis_payload:
            return report.analysis_payload
        return await self.analyze_report(db, report)

    async def get_sources(self, db: Session, report: MedicalReport) -> list[dict[str, Any]]:
        if not report.evidence_sources:
            await self.analyze_report(db, report)
        return report.evidence_sources or []


__all__ = ["MedicalReportAnalysisService", "MedicalTextExtractor", "ReferenceRangeService"]
